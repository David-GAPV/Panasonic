"""
Update IDP Architecture Description — preserves original structure, replaces content only.
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'requirement/IDP_Architecture_Description_Team 7_origin.docx'
OUT = 'requirement/IDP_Architecture_Description_Team 7.docx'

doc = Document(SRC)
paras = doc.paragraphs

def setp(p, text, bold=None, emu=None, hex_color=None, align=None):
    for r in p.runs:
        r.text = ''
    if not p.runs:
        p.add_run()
    r = p.runs[0]
    r.text = text
    if bold is not None: r.bold = bold
    if emu: r.font.size = emu
    if hex_color:
        rv,gv,bv = int(hex_color[:2],16),int(hex_color[2:4],16),int(hex_color[4:],16)
        r.font.color.rgb = RGBColor(rv,gv,bv)
    if align is not None: p.alignment = align

def seth(p, text):
    for r in p.runs:
        r.text = ''
    if not p.runs: p.add_run()
    p.runs[0].text = text

def setb(p, text):
    for r in p.runs:
        r.text = ''
    if not p.runs: p.add_run()
    p.runs[0].text = text

def set_cell(cell, text, bold=False):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ''
        if not p.runs: p.add_run()
        p.runs[0].text = text
        p.runs[0].bold = bold
        p.runs[0].font.size = Pt(10)

C = WD_ALIGN_PARAGRAPH.CENTER

# ── Title page ──────────────────────────────────────────────────────────────
setp(paras[1], 'Panasonic Vietnam', emu=177800, hex_color='666666', align=C)
setp(paras[2], 'Intelligent Document Processing (IDP)', bold=True, emu=304800, hex_color='1F4E79', align=C)
setp(paras[3], 'Architecture Description Document', emu=203200, hex_color='2E75B6', align=C)
setp(paras[5], 'March 2026', emu=152400, hex_color='666666', align=C)
setp(paras[6], 'Team 07', emu=152400, hex_color='666666', align=C)
setp(paras[7], 'Document Version: 2.0 (Updated)', hex_color='999999', align=C)
setp(paras[8], 'Date: March 2026', hex_color='999999', align=C)

# ── Section 2: Conceptual Architecture ──────────────────────────────────────
seth(paras[10], '2. Conceptual Architecture')
seth(paras[11], '')
seth(paras[12], '2.1 Overview')
setp(paras[13],
    'The Conceptual Architecture presents the IDP system as a set of logical layers and services, '
    'independent of any specific technology. It defines what the system does and how its major '
    'components interact, providing a vendor-neutral blueprint. The architecture is organized into '
    'layers from external integrations at the top, through user-facing and backend services, down '
    'to data storage and infrastructure at the base.')
seth(paras[14], '2.2 Layer and Service Descriptions')
setp(paras[15],
    'The following table describes each layer and its constituent services within the Conceptual '
    'Architecture. Each service is described in terms of its responsibility within the overall system.')
seth(paras[18], '2.4 Document Processing Flow')
setp(paras[19], 'The core document processing pipeline follows this sequence:')

# Step titles and bullets (original indices)
STEPS = [
    (20, 'Step 1 - Document Upload', [
        'Users upload document files (DOCX, PDF, PNG, JPEG, TIFF, HEIC, WebP) through the web interface, supporting single and batch uploads.',
        'An optional Shipment Reference is provided to group related documents for cross-document verification.',
        'The system stores the original file in object storage and creates a document record in the database.',
        'The document is immediately queued for classification and extraction.',
    ]),
    (26, 'Step 2 - Document Classification', [
        'The system automatically determines the document type using a score-based classification algorithm.',
        'Scores are computed by scanning extracted text for keywords characteristic of each type: Commercial Invoice, Packing List, Bill of Lading, or Warehouse Receipt.',
        'The highest-scoring type is assigned. Documents that cannot be classified are marked as unknown.',
        'Classification results and confidence scores are attached to the document record.',
    ]),
    (30, 'Step 3 - OCR Extraction', [
        'For image-based documents, the OCR engine extracts raw text. For DOCX files, text is extracted directly from the document structure.',
        'PDF files are converted to images and processed through the OCR engine.',
        'Over 20 structured fields are extracted per document type: invoice number, date, supplier/buyer names, total amount, currency, container numbers, vessel name, ports, bill of lading number, packing list number, warehouse receipt number, total packages, and weights.',
        'Each document receives an overall confidence score based on field completeness.',
    ]),
    (36, 'Step 4 - Intelligent Validation', [
        'Mandatory Field Validation: The system checks that all required fields for the document type are present. Missing fields cause the document to be flagged.',
        'Cross-Document Verification: When a Shipment Reference is provided, 11 overlapping fields are compared across all documents in the same shipment group.',
        'Smart comparison handles container number subsets and unit text normalization.',
        'Validation results are stored and attached to the document record.',
    ]),
    (41, 'Step 5 - Routing Decision', [
        'Documents passing all checks are assigned status "processed" and appear on the main dashboard.',
        'Documents with missing mandatory fields or cross-verification mismatches are flagged as "needs_review" and routed to the Review Queue.',
        'A notification is sent to subscribed operators whenever a document is flagged.',
    ]),
    (45, 'Step 6 - Human Review', [
        'Flagged documents appear in a prioritized Review Queue accessible to Reviewers and Administrators.',
        'Reviewers see the original document side-by-side with extracted data, with validation failures highlighted.',
        'Reviewers can correct field values and approve or reject the document.',
        'All corrections and decisions are recorded in the audit log with reviewer identity and timestamp.',
    ]),
    (51, 'Step 7 - ERP Integration Readiness', [
        'Validated document data is mapped to SAP transaction schemas through the SAP simulation module.',
        'The system simulates SAP MIRO (Invoice Verification) and SAP MIGO (Goods Receipt) transactions using extracted field values.',
        'Extracted data can be exported in CSV or JSON format for downstream system consumption.',
    ]),
    (55, 'Step 8 - Data Export', [
        'Approved document data is available for export in CSV and JSON formats.',
        'PDF and CSV reports summarizing processing volumes, validation statistics, and cross-verification outcomes can be generated on demand.',
    ]),
    (60, 'Step 9 - Notification', [
        'Notifications are sent when a document is flagged due to validation failure or cross-document mismatch.',
        'Subscribed operators receive alerts via the managed notification service, enabling timely response without constant dashboard monitoring.',
    ]),
    (63, 'Step 10 - Audit and Reporting', [
        'Every action is recorded in an immutable audit trail: upload, classification, extraction, validation, cross-verification, review decisions, field edits, and notifications.',
        'Audit records include timestamp, user identity, action type, and descriptive detail.',
        'Processing metrics including volume, confidence distributions, error rates, and validation failure patterns are available through the Analytics module.',
    ]),
]

for title_idx, title_text, bullets in STEPS:
    setb(paras[title_idx], title_text)
    next_titles = [s[0] for s in STEPS if s[0] > title_idx]
    end_idx = next_titles[0] if next_titles else len(paras)
    bps = [p for p in paras[title_idx+1:end_idx]
           if p.style.name in ('whitespace-normal', 'List Paragraph', 'List Bullet')]
    for i, bp in enumerate(bps):
        setb(bp, bullets[i] if i < len(bullets) else '')

# ── Section 3: Solution Architecture ────────────────────────────────────────
seth(paras[68], '3. Solution Architecture (AWS)')
seth(paras[69], '')
seth(paras[70], '3.1 Overview')
setp(paras[71],
    'The Solution Architecture maps each logical service from the Conceptual Architecture to specific '
    'AWS managed services deployed in the Asia Pacific (Singapore) region. The deployed system covers '
    'all core functional requirements: document ingestion, OCR processing, validation, review workflow, '
    'notifications, audit logging, and reporting.')
seth(paras[72], '3.2 Service Mapping')
setp(paras[73],
    'The following table provides the complete mapping from logical services to the AWS implementation '
    'as deployed, including a description of each service\'s role in the running system.')
seth(paras[76], '3.4 Security Architecture')
setp(paras[77],
    'The security architecture implements defense-in-depth across all layers:')

SEC = [
    'Network Security: Amazon VPC with public and private subnet separation. The database resides in a private subnet with no direct internet access. Security groups restrict inbound traffic to necessary protocols only.',
    'Data Encryption: All user-facing traffic is encrypted using TLS. The web application is accessible exclusively over HTTPS through a custom domain with a certificate from a trusted authority.',
    'Identity and Access: Session-based authentication with hashed password storage. Three-role RBAC (Administrator, Reviewer, Uploader) enforced at both server and interface levels.',
    'Audit and Compliance: Application-level audit logs record every user action with timestamp, user identity, and action detail, stored in the relational database.',
    'Least Privilege: A dedicated IAM service role grants compute instances access only to the specific storage, database, and notification resources they require.',
]
sec_bps = [p for p in paras[78:83] if p.style.name == 'List Paragraph']
for i, bp in enumerate(sec_bps):
    if i < len(SEC): setb(bp, SEC[i])

seth(paras[83], '3.5 Scalability and Performance')
setp(paras[84], 'The architecture meets performance targets while supporting future volume growth:')

SCALE = [
    'Compute Separation: The OCR processing service runs on a dedicated compute instance, isolating CPU-intensive processing from the web application to maintain interface responsiveness.',
    'Synchronous Processing: Documents are processed immediately upon upload, providing users with instant feedback on classification, extraction, and validation results.',
    'Database Performance: The managed PostgreSQL database handles document metadata, extracted fields, validation results, and audit logs with query optimization.',
    'Storage Scaling: Amazon S3 provides virtually unlimited, highly durable storage for uploaded document files with automatic scaling.',
]
scale_bps = [p for p in paras[85:89] if p.style.name == 'List Paragraph']
for i, bp in enumerate(scale_bps):
    if i < len(SCALE): setb(bp, SCALE[i])

seth(paras[89], '3.6 Disaster Recovery')
setp(paras[90], 'The disaster recovery strategy leverages managed AWS service capabilities:')

DR = [
    'Database Backups: The managed RDS PostgreSQL instance provides automated daily snapshots and point-in-time recovery.',
    'Object Storage Durability: Amazon S3 provides 99.999999999% (11 nines) durability for uploaded document files through built-in redundancy.',
    'Service Recovery: Application services are configured for automatic restart on failure. Infrastructure provisioning is scripted for rapid re-deployment.',
    'Availability: Current deployment targets single-AZ. The architecture can be extended to multi-AZ with load balancing for higher availability requirements.',
]
dr_bps = [p for p in paras[91:95] if p.style.name == 'List Paragraph']
for i, bp in enumerate(dr_bps):
    if i < len(DR): setb(bp, DR[i])

# ── Section 4: Requirements Traceability ────────────────────────────────────
seth(paras[96], '4. Requirements Traceability')
setp(paras[97],
    'The following summarizes how the deployed architecture addresses the key functional and '
    'non-functional requirements defined in the Business Requirements Document (BRD).')
seth(paras[98], '4.1 Functional Requirements Coverage')

FR = [
    'Document Upload (FR-001 to FR-003): Web application with object storage. Supports DOCX, PDF, PNG, JPEG, TIFF, HEIC, and WebP. Shipment Reference input enables cross-document grouping.',
    'OCR and Extraction (FR-004 to FR-008): Dedicated OCR Service with Tesseract engine. Direct DOCX text extraction. Score-based classification. Over 20 structured fields extracted with confidence scoring.',
    'Validation (FR-009 to FR-012): Mandatory field validation per document type. Cross-document verification across 11 fields within shipment groups. Smart comparison for container numbers and unit normalization.',
    'Review and Workflow (FR-013 to FR-017): Review Queue for flagged documents. Side-by-side view. Field editing, approve, and reject actions for authorized roles.',
    'Enterprise Integration (FR-018 to FR-020): SAP MIRO and MIGO transaction simulation. CSV and JSON export for downstream system consumption.',
    'Audit and Analytics (FR-021 to FR-025): Immutable audit log for all actions. Analytics with processing trends and error analysis. PDF and CSV report generation.',
    'User Management (FR-027): Session-based authentication with hashed passwords. Three-role RBAC enforced at server and interface levels.',
    'Notifications (FR-026, FR-029): Managed notification service delivers alerts when documents are flagged for validation failures or cross-document mismatches.',
]
fr_bps = [p for p in paras[99:107] if p.style.name == 'List Paragraph']
for i, bp in enumerate(fr_bps):
    if i < len(FR): setb(bp, FR[i])

seth(paras[107], '4.2 Non-Functional Requirements Coverage')

NFR = [
    'Performance (NFR-001 to NFR-004): Dedicated OCR compute instance prevents processing load from affecting web interface responsiveness. Synchronous processing provides immediate user feedback.',
    'Scalability (NFR-005, NFR-006): Vertical and horizontal scaling paths for web and OCR instances. Managed database and object storage scale automatically.',
    'Availability (NFR-007, NFR-008): Automatic service restart on failure. Managed database with automated backups. Architecture extensible to multi-AZ deployment.',
    'Security (NFR-011 to NFR-016): VPC isolation, TLS encryption, hashed passwords, RBAC enforcement, IAM least-privilege service role, and private database subnet.',
    'Compliance (NFR-017 to NFR-019): Immutable application-level audit logs recording all user and system actions with timestamps and user identity.',
    'Disaster Recovery (NFR-030, NFR-031): Automated RDS snapshots for database recovery. S3 eleven-nines durability for document files. Scripted infrastructure for rapid re-provisioning.',
    'Cost Efficiency (NFR-040): Right-sized compute instances. Managed services eliminate operational overhead. Pay-per-use object storage and notification services.',
]
nfr_bps = [p for p in paras[108:115] if p.style.name == 'List Paragraph']
for i, bp in enumerate(nfr_bps):
    if i < len(NFR): setb(bp, NFR[i])

# ── Table 0: Conceptual Architecture (only actually deployed services) ────────
T0_ROWS = [
    ('External Services', "SAP's APIs",
     'ERP integration for invoice verification and goods receipt posting. Demonstrated through the SAP MIRO/MIGO simulation module in the current deployment.'),
    ('User Facing', 'User Interface / Front-end',
     'Web-based application for document upload, review, and dashboard access. Served over HTTPS with role-aware navigation for three user roles.'),
    ('Integration Layer', 'DNS / TLS',
     'Custom domain with TLS certificate routing users to the web application over HTTPS.'),
    ('', 'Reverse Proxy',
     'Handles TLS termination, HTTPS enforcement, and request forwarding to the web application service.'),
    ('', 'Data Export',
     'CSV and JSON export providing structured document data for downstream system integration.'),
    ('Backend Services', 'Review and Workflow',
     'Manages the review queue, document approval workflow, field editing, and approve/reject actions for authorized reviewers.'),
    ('', 'User/Auth Service',
     'Session-based authentication and role-based access control for Administrator, Reviewer, and Uploader roles.'),
    ('', 'Document Service',
     'Core document lifecycle: upload, storage, classification, field extraction, validation, and cross-document verification.'),
    ('', 'Analytical and Reporting',
     'Generates processing statistics, error analytics, and downloadable PDF/CSV reports.'),
    ('', 'Notification Service',
     'Delivers alerts to subscribed operators when documents are flagged for validation failures or cross-document mismatches.'),
    ('OCR Core Services', 'OCR Service',
     'Text extraction from image-based documents (PNG, JPEG, TIFF, HEIC, WebP) and direct extraction from DOCX files. Score-based document type classification.'),
    ('', 'Validation Service',
     'Mandatory field validation per document type and cross-document consistency checks across 11 fields within shipment groups.'),
    ('Data Layer', 'Application Database',
     'Relational database storing document metadata, extracted fields, validation results, cross-verification outcomes, user accounts, and audit logs.'),
    ('Storage Layer', 'Object Storage',
     'Durable storage for uploaded document files, referenced by unique identifiers in document metadata records.'),
    ('Networking and Security', 'Authentication and Authorization',
     'Session-based authentication with hashed passwords. Three-role RBAC enforced at server and interface levels.'),
    ('', 'Network Security',
     'VPC with public/private subnet separation. Security groups restrict inbound traffic. Database isolated in private subnet.'),
    ('', 'IAM',
     'Service role granting compute instances least-privilege access to storage, database, and notification services.'),
    ('Monitoring', 'Audit Log',
     'Immutable record of all user and system actions with timestamp, user identity, and action detail.'),
    ('', 'Dashboard and Analytics',
     'Real-time KPI dashboard and analytics module for processing volume, confidence distributions, and error rate trends.'),
]

# ── Table 1: AWS Service Mapping (only actually deployed services) ────────────
T1_ROWS = [
    ('User Facing', 'User Interface / Front-end', 'Amazon EC2 (Web Server)',
     'Flask web application on EC2 serves the full user interface: dashboard, upload, review queue, document detail, analytics, reports, and SAP simulation.'),
    ('Integration Layer', 'DNS / TLS', 'Custom Domain + ACM/Certbot',
     'Custom domain with TLS certificate providing HTTPS access. Certificate provisioned and auto-renewed via Certbot.'),
    ('', 'Reverse Proxy', 'Nginx on EC2',
     'Nginx on the Web Server EC2 instance handles TLS termination, HTTPS enforcement, and request forwarding to the Flask application.'),
    ('', 'Data Export', 'Web Application (EC2)',
     'CSV and JSON export endpoints within the Flask application provide structured data for downstream system integration.'),
    ('Backend Services', 'Review and Workflow', 'Amazon EC2 (Web Server)',
     'Review queue, side-by-side document view, field editing, and approve/reject workflow within the Flask application.'),
    ('', 'User/Auth Service', 'Amazon EC2 (Web Server)',
     'Session-based authentication and three-role RBAC (Administrator, Reviewer, Uploader) within the Flask application.'),
    ('', 'Document Service', 'Amazon EC2 (Web Server)',
     'Document lifecycle management: upload handling, OCR service orchestration, validation, cross-verification, and status management.'),
    ('', 'Analytical and Reporting', 'Amazon EC2 (Web Server)',
     'Analytics and reporting module generating processing statistics, error analysis, and downloadable PDF/CSV reports.'),
    ('', 'Notification Service', 'Amazon SNS',
     'Managed publish-subscribe service delivering email alerts to subscribed operators when documents are flagged.'),
    ('OCR Core Services', 'OCR Service', 'Amazon EC2 (OCR Server)',
     'Dedicated EC2 instance running the Tesseract OCR engine. Handles image-based extraction and direct DOCX text parsing. Performs score-based document classification.'),
    ('', 'Validation Service', 'Amazon EC2 (Web Server)',
     'Mandatory field validation and cross-document verification triggered automatically after each document is processed.'),
    ('Data Layer', 'Application Database', 'Amazon RDS (PostgreSQL)',
     'Managed PostgreSQL database storing document metadata, extracted fields, validation results, cross-verification outcomes, user accounts, and audit logs. Automated daily backups enabled.'),
    ('Storage Layer', 'Object Storage', 'Amazon S3',
     'Scalable object storage for uploaded document files. Files stored with unique keys and retrieved on demand for viewing or download.'),
    ('Networking and Security', 'Network Isolation', 'Amazon VPC',
     'Virtual Private Cloud with public subnets for EC2 instances and a private subnet for RDS. Security groups control inbound and outbound traffic.'),
    ('', 'IAM Service Role', 'AWS IAM',
     'Service role granting EC2 instances least-privilege access to S3, RDS, and SNS only.'),
    ('Monitoring', 'Audit Log', 'Amazon RDS (PostgreSQL)',
     'All user and system actions recorded in the audit log table within the PostgreSQL database, surfaced on document detail pages and in reports.'),
    ('', 'Dashboard and Analytics', 'Amazon EC2 (Web Server)',
     'Built-in web dashboard with real-time KPI cards and analytics module for processing volume, confidence distributions, and error rate trends.'),
]

def rebuild_table(table, rows):
    needed = 1 + len(rows)
    while len(table.rows) < needed:
        table.add_row()
    while len(table.rows) > needed:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            if c_idx < len(table.columns):
                set_cell(table.rows[r_idx + 1].cells[c_idx], str(val))

rebuild_table(doc.tables[0], T0_ROWS)
rebuild_table(doc.tables[1], T1_ROWS)

doc.save(OUT)
sys.stdout.write(f'Saved: {OUT}\n')
sys.stdout.flush()
