"""Create test DOCX documents for cross-document verification testing.
Set A: 4 documents with CONSISTENT data (same shipment) — should pass cross-verify.
Set B: 4 documents with DELIBERATE MISMATCHES — should trigger cross-verify failures.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(45, 66, 133)

def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_v = p.add_run(value)
    run_v.font.size = Pt(11)

# ============================================================
# SHARED SHIPMENT DATA (consistent across Set A)
# ============================================================
SHIPMENT = {
    "po_number": "PO-2025-TEST-001",
    "invoice_number": "INV-2025-TEST-7001",
    "bl_number": "HDMU1234567890",
    "supplier": "Shenzhen Greentech Electronics Co., Ltd.",
    "buyer": "Panasonic Appliances Vietnam Co., Ltd.",
    "vessel": "EVER GOLDEN V.025E",
    "port_of_loading": "Shenzhen Yantian Port",
    "port_of_discharge": "Hai Phong Port",
    "container_no": "HDMU7654321",
    "gross_weight": "12,500.00 kg",
    "net_weight": "11,800.00 kg",
    "total_packages": "200",
    "total_amount": "85,000.00",
    "currency": "USD",
    "date": "10 March 2025",
    "tax_code": "0101248141",
    "incoterms": "CIF Hai Phong",
    "seal_no": "SL-20250310-A1",
}

# ============================================================
# SET A: Consistent documents (same shipment, matching data)
# ============================================================

def create_invoice_a():
    doc = Document()
    add_heading(doc, "COMMERCIAL INVOICE", level=1)
    doc.add_paragraph("")
    add_kv(doc, "Invoice No.", SHIPMENT["invoice_number"])
    add_kv(doc, "Date", SHIPMENT["date"])
    add_kv(doc, "PO No.", SHIPMENT["po_number"])
    add_kv(doc, "Currency", SHIPMENT["currency"])
    add_kv(doc, "Incoterms", SHIPMENT["incoterms"])
    doc.add_paragraph("")
    add_heading(doc, "Seller", level=2)
    doc.add_paragraph(SHIPMENT["supplier"])
    add_heading(doc, "Buyer", level=2)
    doc.add_paragraph(SHIPMENT["buyer"])
    doc.add_paragraph(f"Tax Code (MST): {SHIPMENT['tax_code']}")
    doc.add_paragraph("")
    add_heading(doc, "Shipment Details", level=2)
    add_kv(doc, "Vessel / Voyage", SHIPMENT["vessel"])
    add_kv(doc, "Port of Loading", SHIPMENT["port_of_loading"])
    add_kv(doc, "Port of Discharge", SHIPMENT["port_of_discharge"])
    add_kv(doc, "B/L No.", SHIPMENT["bl_number"])
    add_kv(doc, "Container No", SHIPMENT["container_no"])
    doc.add_paragraph("")
    add_heading(doc, "Items", level=2)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    headers = ["No.", "Description", "Qty", "Unit Price", "Amount"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1", "LED Panel Module LP-500", "500", "80.00", "40,000.00"],
        ["2", "Power Supply Unit PSU-300", "300", "100.00", "30,000.00"],
        ["3", "Control Board CB-200", "200", "75.00", "15,000.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")
    doc.add_paragraph(f"TOTAL CIF VALUE: USD {SHIPMENT['total_amount']}")
    doc.add_paragraph(f"Total Packages: {SHIPMENT['total_packages']} cartons")
    doc.add_paragraph(f"Gross Weight: {SHIPMENT['gross_weight']}")
    doc.add_paragraph(f"Net Weight: {SHIPMENT['net_weight']}")
    doc.save("doc_input/TestA_01_Invoice_INV-2025-TEST-7001.docx")
    print("Created: TestA_01_Invoice")

def create_packing_list_a():
    doc = Document()
    add_heading(doc, "PACKING LIST", level=1)
    doc.add_paragraph("")
    add_kv(doc, "Packing List No.", "PL-2025-TEST-7001")
    add_kv(doc, "Date", SHIPMENT["date"])
    add_kv(doc, "Invoice No.", SHIPMENT["invoice_number"])
    add_kv(doc, "PO No.", SHIPMENT["po_number"])
    doc.add_paragraph("")
    add_heading(doc, "Shipper", level=2)
    doc.add_paragraph(SHIPMENT["supplier"])
    add_heading(doc, "Consignee", level=2)
    doc.add_paragraph(SHIPMENT["buyer"])
    doc.add_paragraph("")
    add_heading(doc, "Packing Details", level=2)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    headers = ["Carton No", "Description", "Qty/Ctn", "N.W. (kg)", "G.W. (kg)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1-100", "LED Panel Module LP-500", "5", "5,900.00", "6,250.00"],
        ["101-160", "Power Supply Unit PSU-300", "5", "3,540.00", "3,750.00"],
        ["161-200", "Control Board CB-200", "5", "2,360.00", "2,500.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")
    doc.add_paragraph(f"Total Packages: {SHIPMENT['total_packages']} cartons")
    doc.add_paragraph(f"Gross Weight: {SHIPMENT['gross_weight']}")
    doc.add_paragraph(f"Net Weight: {SHIPMENT['net_weight']}")
    doc.add_paragraph(f"Container No: {SHIPMENT['container_no']}")
    doc.add_paragraph(f"Seal No.: {SHIPMENT['seal_no']}")
    doc.add_paragraph(f"Vessel / Voyage: {SHIPMENT['vessel']}")
    doc.add_paragraph(f"Port of Loading: {SHIPMENT['port_of_loading']}")
    doc.add_paragraph(f"Port of Discharge: {SHIPMENT['port_of_discharge']}")
    doc.save("doc_input/TestA_02_PackingList_PL-2025-TEST-7001.docx")
    print("Created: TestA_02_PackingList")

def create_bl_a():
    doc = Document()
    add_heading(doc, "BILL OF LADING", level=1)
    doc.add_paragraph("")
    add_kv(doc, "B/L No.", SHIPMENT["bl_number"])
    add_kv(doc, "Booking Ref", "BK-20250308-YT")
    add_kv(doc, "Date", SHIPMENT["date"])
    doc.add_paragraph("")
    add_heading(doc, "Shipper", level=2)
    doc.add_paragraph(SHIPMENT["supplier"])
    add_heading(doc, "Consignee", level=2)
    doc.add_paragraph(SHIPMENT["buyer"])
    doc.add_paragraph("")
    add_heading(doc, "Voyage Details", level=2)
    add_kv(doc, "Vessel / Voyage", SHIPMENT["vessel"])
    add_kv(doc, "Port of Loading", SHIPMENT["port_of_loading"])
    add_kv(doc, "Port of Discharge", SHIPMENT["port_of_discharge"])
    add_kv(doc, "ETD", "08 March 2025")
    add_kv(doc, "ETA", "12 March 2025")
    doc.add_paragraph("")
    add_heading(doc, "Cargo Details", level=2)
    add_kv(doc, "Container No", SHIPMENT["container_no"])
    add_kv(doc, "Seal No.", SHIPMENT["seal_no"])
    add_kv(doc, "Total Packages", f"{SHIPMENT['total_packages']} cartons")
    add_kv(doc, "Gross Weight", SHIPMENT["gross_weight"])
    add_kv(doc, "Net Weight", SHIPMENT["net_weight"])
    add_kv(doc, "Measurement", "45.50 CBM")
    doc.add_paragraph("")
    doc.add_paragraph("Description: Electronic components — LED panels, power supplies, control boards")
    doc.add_paragraph(f"Invoice No.: {SHIPMENT['invoice_number']}")
    doc.add_paragraph(f"PO No.: {SHIPMENT['po_number']}")
    doc.add_paragraph(f"Ocean Freight: USD 2,800.00")
    doc.save("doc_input/TestA_03_BillOfLading_HDMU1234567890.docx")
    print("Created: TestA_03_BillOfLading")

def create_warehouse_receipt_a():
    doc = Document()
    add_heading(doc, "WAREHOUSE RECEIPT", level=1)
    doc.add_paragraph("")
    add_kv(doc, "WH Receipt No.", "WR-2025-HP-00101")
    add_kv(doc, "Date", "12 March 2025")
    add_kv(doc, "Date Received", "12 March 2025")
    doc.add_paragraph("")
    add_heading(doc, "Shipment Reference", level=2)
    add_kv(doc, "Invoice No.", SHIPMENT["invoice_number"])
    add_kv(doc, "PO No.", SHIPMENT["po_number"])
    add_kv(doc, "B/L No.", SHIPMENT["bl_number"])
    add_kv(doc, "Container No", SHIPMENT["container_no"])
    doc.add_paragraph("")
    add_heading(doc, "Supplier", level=2)
    doc.add_paragraph(SHIPMENT["supplier"])
    add_heading(doc, "Receiver", level=2)
    doc.add_paragraph(SHIPMENT["buyer"])
    doc.add_paragraph("")
    add_heading(doc, "Goods Received", level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ["Description", "Expected Qty", "Received Qty", "Condition"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["LED Panel Module LP-500", "500", "500", "Good"],
        ["Power Supply Unit PSU-300", "300", "300", "Good"],
        ["Control Board CB-200", "200", "200", "Good"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")
    doc.add_paragraph(f"Total Packages: {SHIPMENT['total_packages']} cartons")
    doc.add_paragraph(f"Gross Weight: {SHIPMENT['gross_weight']}")
    doc.add_paragraph(f"Customs Clearance: 12 March 2025")
    doc.add_paragraph(f"Port of Discharge: {SHIPMENT['port_of_discharge']}")
    doc.add_paragraph(f"Vessel / Voyage: {SHIPMENT['vessel']}")
    doc.save("doc_input/TestA_04_WarehouseReceipt_WR-2025-HP-00101.docx")
    print("Created: TestA_04_WarehouseReceipt")

# ============================================================
# SET B: Documents with DELIBERATE MISMATCHES (same PO, different values)
# ============================================================

def create_invoice_b():
    """Invoice with CORRECT data — baseline for Set B."""
    doc = Document()
    add_heading(doc, "COMMERCIAL INVOICE", level=1)
    doc.add_paragraph("")
    add_kv(doc, "Invoice No.", "INV-2025-TEST-8002")
    add_kv(doc, "Date", "15 March 2025")
    add_kv(doc, "PO No.", "PO-2025-TEST-002")
    add_kv(doc, "Currency", "USD")
    add_kv(doc, "Incoterms", "CIF Hai Phong")
    doc.add_paragraph("")
    add_heading(doc, "Seller", level=2)
    doc.add_paragraph("Shenzhen Greentech Electronics Co., Ltd.")
    add_heading(doc, "Buyer", level=2)
    doc.add_paragraph("Panasonic Appliances Vietnam Co., Ltd.")
    doc.add_paragraph("Tax Code (MST): 0101248141")
    doc.add_paragraph("")
    add_heading(doc, "Shipment Details", level=2)
    add_kv(doc, "Vessel / Voyage", "COSCO STAR V.018W")
    add_kv(doc, "Port of Loading", "Shanghai Yangshan Port")
    add_kv(doc, "Port of Discharge", "Hai Phong Port")
    add_kv(doc, "B/L No.", "COSU9876543210")
    add_kv(doc, "Container No", "COSU1122334")
    doc.add_paragraph("")
    doc.add_paragraph("TOTAL CIF VALUE: USD 120,000.00")
    doc.add_paragraph("Total Packages: 350 cartons")
    doc.add_paragraph("Gross Weight: 18,200.00 kg")
    doc.add_paragraph("Net Weight: 17,100.00 kg")
    doc.save("doc_input/TestB_01_Invoice_INV-2025-TEST-8002.docx")
    print("Created: TestB_01_Invoice (baseline)")

def create_packing_list_b():
    """Packing list with MISMATCHES: different gross weight, net weight, total packages."""
    doc = Document()
    add_heading(doc, "PACKING LIST", level=1)
    doc.add_paragraph("")
    add_kv(doc, "Packing List No.", "PL-2025-TEST-8002")
    add_kv(doc, "Date", "15 March 2025")
    add_kv(doc, "Invoice No.", "INV-2025-TEST-8002")  # SAME invoice — links to Set B
    add_kv(doc, "PO No.", "PO-2025-TEST-002")         # SAME PO — links to Set B
    doc.add_paragraph("")
    add_heading(doc, "Shipper", level=2)
    doc.add_paragraph("Shenzhen Greentech Electronics Co., Ltd.")
    add_heading(doc, "Consignee", level=2)
    doc.add_paragraph("Panasonic Appliances Vietnam Co., Ltd.")
    doc.add_paragraph("")
    doc.add_paragraph("Total Packages: 380 cartons")           # MISMATCH: 380 vs 350
    doc.add_paragraph("Gross Weight: 19,500.00 kg")            # MISMATCH: 19,500 vs 18,200
    doc.add_paragraph("Net Weight: 17,100.00 kg")              # Same
    doc.add_paragraph("Container No: COSU1122334")
    doc.add_paragraph(f"Vessel / Voyage: COSCO STAR V.018W")
    doc.add_paragraph(f"Port of Loading: Shanghai Yangshan Port")
    doc.add_paragraph(f"Port of Discharge: Hai Phong Port")
    doc.save("doc_input/TestB_02_PackingList_PL-2025-TEST-8002.docx")
    print("Created: TestB_02_PackingList (MISMATCHES: gross_weight, total_packages)")

def create_bl_b():
    """Bill of Lading with MISMATCHES: different vessel, port of loading."""
    doc = Document()
    add_heading(doc, "BILL OF LADING", level=1)
    doc.add_paragraph("")
    add_kv(doc, "B/L No.", "COSU9876543210")           # SAME B/L — links to Set B
    add_kv(doc, "Booking Ref", "BK-20250313-YS")
    add_kv(doc, "Date", "15 March 2025")
    doc.add_paragraph("")
    add_heading(doc, "Shipper", level=2)
    doc.add_paragraph("Shenzhen Greentech Electronics Co., Ltd.")
    add_heading(doc, "Consignee", level=2)
    doc.add_paragraph("Panasonic Appliances Vietnam Co., Ltd.")
    doc.add_paragraph("")
    add_heading(doc, "Voyage Details", level=2)
    add_kv(doc, "Vessel / Voyage", "EVER FORTUNE V.033E")       # MISMATCH: different vessel
    add_kv(doc, "Port of Loading", "Ningbo Beilun Port")        # MISMATCH: different port
    add_kv(doc, "Port of Discharge", "Hai Phong Port")          # Same
    add_kv(doc, "ETD", "13 March 2025")
    add_kv(doc, "ETA", "18 March 2025")
    doc.add_paragraph("")
    add_heading(doc, "Cargo Details", level=2)
    add_kv(doc, "Container No", "COSU1122334")
    add_kv(doc, "Total Packages", "350 cartons")
    add_kv(doc, "Gross Weight", "18,200.00 kg")
    add_kv(doc, "Net Weight", "17,100.00 kg")
    doc.add_paragraph("")
    doc.add_paragraph(f"Invoice No.: INV-2025-TEST-8002")       # Links to Set B
    doc.add_paragraph(f"PO No.: PO-2025-TEST-002")
    doc.add_paragraph(f"Ocean Freight: USD 3,200.00")
    doc.save("doc_input/TestB_03_BillOfLading_COSU9876543210.docx")
    print("Created: TestB_03_BillOfLading (MISMATCHES: vessel, port_of_loading)")

def create_warehouse_receipt_b():
    """Warehouse receipt with MISMATCHES: different supplier name, port of discharge."""
    doc = Document()
    add_heading(doc, "WAREHOUSE RECEIPT", level=1)
    doc.add_paragraph("")
    add_kv(doc, "WH Receipt No.", "WR-2025-HP-00202")
    add_kv(doc, "Date", "18 March 2025")
    add_kv(doc, "Date Received", "18 March 2025")
    doc.add_paragraph("")
    add_heading(doc, "Shipment Reference", level=2)
    add_kv(doc, "Invoice No.", "INV-2025-TEST-8002")   # Links to Set B
    add_kv(doc, "PO No.", "PO-2025-TEST-002")
    add_kv(doc, "B/L No.", "COSU9876543210")
    add_kv(doc, "Container No", "COSU1122334")
    doc.add_paragraph("")
    add_heading(doc, "Supplier", level=2)
    doc.add_paragraph("Shanghai Greentech Electronics Co., Ltd.")  # MISMATCH: Shanghai vs Shenzhen
    add_heading(doc, "Receiver", level=2)
    doc.add_paragraph("Panasonic Appliances Vietnam Co., Ltd.")
    doc.add_paragraph("")
    add_heading(doc, "Goods Received", level=2)
    doc.add_paragraph("Total Packages: 350 cartons")
    doc.add_paragraph("Gross Weight: 18,200.00 kg")
    doc.add_paragraph("Customs Clearance: 18 March 2025")
    doc.add_paragraph("Port of Discharge: Da Nang Port")           # MISMATCH: Da Nang vs Hai Phong
    doc.add_paragraph(f"Vessel / Voyage: COSCO STAR V.018W")
    doc.save("doc_input/TestB_04_WarehouseReceipt_WR-2025-HP-00202.docx")
    print("Created: TestB_04_WarehouseReceipt (MISMATCHES: supplier_name, port_of_discharge)")

# ============================================================
# Generate all
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("SET A: Consistent shipment (should PASS cross-verify)")
    print("=" * 60)
    create_invoice_a()
    create_packing_list_a()
    create_bl_a()
    create_warehouse_receipt_a()
    print()
    print("=" * 60)
    print("SET B: Deliberate mismatches (should FAIL cross-verify)")
    print("=" * 60)
    create_invoice_b()
    create_packing_list_b()
    create_bl_b()
    create_warehouse_receipt_b()
    print()
    print("UPLOAD ORDER:")
    print("  Set A: Upload all 4 files together — expect 0 cross-verify failures")
    print("         (first doc has no related docs yet, subsequent ones match)")
    print("  Set B: Upload all 4 files one by one — expect cross-verify failures:")
    print("         - TestB_02 vs TestB_01: gross_weight, total_packages mismatch")
    print("         - TestB_03 vs TestB_01/02: vessel, port_of_loading mismatch")
    print("         - TestB_04 vs TestB_01/02/03: supplier_name, port_of_discharge mismatch")
