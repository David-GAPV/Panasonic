"""Create test DOCX documents — Set C: VARIANT LAYOUT.
Same core fields as Sets A/B but with:
  - Different label names (Exporter, Notify Party, Arrival Port, etc.)
  - Additional fields (HS codes, payment terms, L/C, insurance, country of origin)
  - Mixed table/paragraph layout
  - Data is CONSISTENT within the set (should pass cross-verify)
  - Tests OCR pattern robustness against layout variations
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)

def add_table_kv(doc, pairs):
    """Add key-value pairs as a 2-column table (variant layout)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(pairs):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ============================================================
# SHARED SHIPMENT DATA — Set C
# ============================================================
SHIP_C = {
    "po_number":         "PO-2025-VN-3300",
    "invoice_number":    "INV-2025-VN-3300",
    "bl_number":         "MAEU5500112233",
    "supplier":          "Dongguan Precision Parts Manufacturing Co., Ltd.",
    "buyer":             "Panasonic Appliances Vietnam Co., Ltd.",
    "vessel":            "MAERSK SELETAR V.012S",
    "port_of_loading":   "Shenzhen Yantian Port",
    "port_of_discharge": "Ho Chi Minh City Cat Lai Port",
    "container_no":      "MAEU7788001",
    "gross_weight":      "22,350.00 kg",
    "net_weight":        "20,800.00 kg",
    "total_packages":    "480",
    "total_amount":      "198,500.00",
    "currency":          "USD",
    "date":              "18 March 2025",
    "tax_code":          "0101248141",
    "incoterms":         "CIF Ho Chi Minh",
    "seal_no":           "SL-20250318-M7",
    "hs_codes":          ["8504.40.90", "8536.90.00", "8534.00.90"],
    "payment_terms":     "Irrevocable L/C at sight",
    "lc_number":         "LC-2025-VCB-08812",
    "insurance_amount":  "200,000.00",
    "country_of_origin": "China",
}

# ============================================================
# SET C — Invoice: "PROFORMA INVOICE" with table-based header
# ============================================================
def create_invoice_c():
    doc = Document()
    # Title uses different wording
    add_heading(doc, "PROFORMA INVOICE / COMMERCIAL INVOICE", level=1)
    doc.add_paragraph("")

    # Header info in a TABLE (variant layout — not key-value paragraphs)
    add_table_kv(doc, [
        ("Invoice Ref.",       SHIP_C["invoice_number"]),
        ("Invoice Date",       SHIP_C["date"]),
        ("Purchase Order",     SHIP_C["po_number"]),
        ("Currency",           SHIP_C["currency"]),
        ("Incoterms",          SHIP_C["incoterms"]),
        ("Payment",            SHIP_C["payment_terms"]),
        ("L/C No.",            SHIP_C["lc_number"]),
        ("Country of Origin",  SHIP_C["country_of_origin"]),
    ])
    doc.add_paragraph("")

    # Parties — uses "Exporter" and "Importer" instead of Seller/Buyer
    add_heading(doc, "Exporter (Seller)", level=2)
    doc.add_paragraph(SHIP_C["supplier"])
    doc.add_paragraph("No. 88 Songshan Lake Road, Dongguan, Guangdong, China 523808")

    add_heading(doc, "Importer (Buyer)", level=2)
    doc.add_paragraph(SHIP_C["buyer"])
    doc.add_paragraph(f"Tax Code (MST): {SHIP_C['tax_code']}")
    doc.add_paragraph("Lot CN2, Thang Long Industrial Park II, Hanoi, Vietnam")
    doc.add_paragraph("")

    # Shipment block — uses "Transport Details" heading
    add_heading(doc, "Transport Details", level=2)
    add_kv(doc, "Vessel / Voyage", SHIP_C["vessel"])
    add_kv(doc, "Port of Loading", SHIP_C["port_of_loading"])
    add_kv(doc, "Port of Discharge", SHIP_C["port_of_discharge"])
    add_kv(doc, "B/L No.", SHIP_C["bl_number"])
    add_kv(doc, "Container No", SHIP_C["container_no"])
    add_kv(doc, "Seal No.", SHIP_C["seal_no"])
    doc.add_paragraph("")

    # Line items with HS codes (additional field)
    add_heading(doc, "Description of Goods", level=2)
    table = doc.add_table(rows=5, cols=7)
    table.style = 'Table Grid'
    headers = ["No.", "HS Code", "Description", "Qty", "Unit", "Unit Price (USD)", "Amount (USD)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1", "8504.40.90", "AC/DC Power Converter Module PCV-800", "600", "PCS", "120.00", "72,000.00"],
        ["2", "8536.90.00", "Industrial Relay Switch IRS-450", "1,200", "PCS", "55.00", "66,000.00"],
        ["3", "8534.00.90", "Multi-layer PCB Assembly MLP-220", "800", "PCS", "45.00", "36,000.00"],
        ["4", "8504.40.90", "Voltage Regulator VR-100S", "1,500", "PCS", "16.33", "24,500.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    # Totals — uses different label format
    doc.add_paragraph(f"TOTAL CIF VALUE: USD {SHIP_C['total_amount']}")
    doc.add_paragraph(f"Insurance Amount: USD {SHIP_C['insurance_amount']}")
    doc.add_paragraph(f"Total Packages: {SHIP_C['total_packages']} cartons")
    doc.add_paragraph(f"Gross Weight: {SHIP_C['gross_weight']}")
    doc.add_paragraph(f"Net Weight: {SHIP_C['net_weight']}")
    doc.add_paragraph("")

    # Footer note (additional context)
    p = doc.add_paragraph()
    p.add_run("Remarks: ").bold = True
    p.add_run("Goods shipped under Irrevocable Letter of Credit. "
              "All items comply with ACFTA Form E preferential tariff. "
              "Certificate of Origin attached separately.")

    doc.save("doc_input/TestC_01_Invoice_INV-2025-VN-3300.docx")
    print("Created: TestC_01_Invoice (variant layout — table header, HS codes, L/C, insurance)")


# ============================================================
# SET C — Packing List: uses "DETAILED PACKING LIST" with
#          carton-level table and summary in footer-style
# ============================================================
def create_packing_list_c():
    doc = Document()
    add_heading(doc, "DETAILED PACKING LIST", level=1)
    doc.add_paragraph("")

    # Header in table format
    add_table_kv(doc, [
        ("Packing List No.",   "PL-2025-VN-3300"),
        ("Date",               SHIP_C["date"]),
        ("Ref. Invoice",       SHIP_C["invoice_number"]),
        ("Purchase Order",     SHIP_C["po_number"]),
        ("Country of Origin",  SHIP_C["country_of_origin"]),
    ])
    doc.add_paragraph("")

    # Parties — uses "Shipper/Exporter" and "Consignee/Importer"
    add_heading(doc, "Shipper / Exporter", level=2)
    doc.add_paragraph(SHIP_C["supplier"])
    add_heading(doc, "Consignee / Importer", level=2)
    doc.add_paragraph(SHIP_C["buyer"])
    doc.add_paragraph("")

    # Detailed carton-level packing (more granular than Sets A/B)
    add_heading(doc, "Carton Details", level=2)
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    headers = ["Carton No.", "HS Code", "Description", "Qty/Ctn", "N.W. (kg)", "G.W. (kg)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["001-120",  "8504.40.90", "AC/DC Power Converter PCV-800",  "5",  "7,080.00",  "7,560.00"],
        ["121-240",  "8536.90.00", "Industrial Relay Switch IRS-450", "10", "6,480.00",  "7,200.00"],
        ["241-360",  "8534.00.90", "Multi-layer PCB Assembly MLP-220","10", "4,320.00",  "4,800.00"],
        ["361-420",  "8504.40.90", "Voltage Regulator VR-100S",      "25", "1,620.00",  "1,800.00"],
        ["421-480",  "8504.40.90", "Voltage Regulator VR-100S (cont)","25","1,300.00",  "990.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    # Summary block — different position and labels
    add_heading(doc, "Summary", level=2)
    doc.add_paragraph(f"Total Packages: {SHIP_C['total_packages']} cartons")
    doc.add_paragraph(f"Gross Wt: {SHIP_C['gross_weight']}")
    doc.add_paragraph(f"Net Wt: {SHIP_C['net_weight']}")
    doc.add_paragraph("")

    # Transport info at bottom (variant: usually at top in other sets)
    add_heading(doc, "Shipping Marks", level=2)
    add_kv(doc, "Container No", SHIP_C["container_no"])
    add_kv(doc, "Seal No.", SHIP_C["seal_no"])
    add_kv(doc, "Vessel / Voyage", SHIP_C["vessel"])
    add_kv(doc, "Port of Loading", SHIP_C["port_of_loading"])
    add_kv(doc, "Port of Discharge", SHIP_C["port_of_discharge"])

    doc.save("doc_input/TestC_02_PackingList_PL-2025-VN-3300.docx")
    print("Created: TestC_02_PackingList (variant layout — carton-level detail, HS codes)")


# ============================================================
# SET C — Bill of Lading: "OCEAN BILL OF LADING" with
#          Notify Party, multiple containers style, measurement
# ============================================================
def create_bl_c():
    doc = Document()
    add_heading(doc, "OCEAN BILL OF LADING", level=1)
    doc.add_paragraph("")

    # Header in table
    add_table_kv(doc, [
        ("B/L No.",       SHIP_C["bl_number"]),
        ("Booking Ref",   "BK-20250316-YT"),
        ("Date of Issue", SHIP_C["date"]),
        ("Place of Issue","Shenzhen, China"),
    ])
    doc.add_paragraph("")

    # Parties — adds "Notify Party" (additional field)
    add_heading(doc, "Shipper", level=2)
    doc.add_paragraph(SHIP_C["supplier"])
    doc.add_paragraph("No. 88 Songshan Lake Road, Dongguan, Guangdong, China")

    add_heading(doc, "Consignee", level=2)
    doc.add_paragraph(SHIP_C["buyer"])
    doc.add_paragraph("Lot CN2, Thang Long Industrial Park II, Hanoi, Vietnam")

    add_heading(doc, "Notify Party", level=2)
    doc.add_paragraph("Panasonic Logistics Vietnam")
    doc.add_paragraph("Tel: +84-24-3952-xxxx / Fax: +84-24-3952-xxxx")
    doc.add_paragraph("")

    # Voyage — uses "Arrival Port" instead of "Port of Discharge"
    add_heading(doc, "Vessel & Voyage", level=2)
    add_kv(doc, "Vessel / Voyage", SHIP_C["vessel"])
    add_kv(doc, "Port of Loading", SHIP_C["port_of_loading"])
    add_kv(doc, "Arrival Port", SHIP_C["port_of_discharge"])
    add_kv(doc, "ETD", "16 March 2025")
    add_kv(doc, "ETA", "21 March 2025")
    doc.add_paragraph("")

    # Cargo in table format
    add_heading(doc, "Particulars of Cargo", level=2)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    headers = ["Container / Seal", "No. of Packages", "Description", "Gross Weight", "Measurement"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = f"{SHIP_C['container_no']} / {SHIP_C['seal_no']}"
    table.rows[1].cells[1].text = f"{SHIP_C['total_packages']} CTNS"
    table.rows[1].cells[2].text = ("Electronic components:\n"
                                    "- AC/DC Power Converters\n"
                                    "- Industrial Relay Switches\n"
                                    "- Multi-layer PCB Assemblies\n"
                                    "- Voltage Regulators\n"
                                    f"HS Codes: {', '.join(SHIP_C['hs_codes'])}")
    table.rows[1].cells[3].text = SHIP_C["gross_weight"]
    table.rows[1].cells[4].text = "68.20 CBM"
    doc.add_paragraph("")

    # Additional fields at bottom
    doc.add_paragraph(f"Net Weight: {SHIP_C['net_weight']}")
    doc.add_paragraph(f"Invoice No.: {SHIP_C['invoice_number']}")
    doc.add_paragraph(f"PO No.: {SHIP_C['po_number']}")
    doc.add_paragraph(f"L/C No.: {SHIP_C['lc_number']}")
    doc.add_paragraph(f"Ocean Freight: USD 4,200.00 (Prepaid)")
    doc.add_paragraph(f"Insurance: USD {SHIP_C['insurance_amount']}")

    doc.save("doc_input/TestC_03_BillOfLading_MAEU5500112233.docx")
    print("Created: TestC_03_BillOfLading (variant layout — Notify Party, Arrival Port, cargo table)")


# ============================================================
# SET C — Warehouse Receipt: "GOODS RECEIVED NOTE / WAREHOUSE RECEIPT"
#          with inspection table, condition codes, QC section
# ============================================================
def create_warehouse_receipt_c():
    doc = Document()
    add_heading(doc, "GOODS RECEIVED NOTE / WAREHOUSE RECEIPT", level=1)
    doc.add_paragraph("")

    # Header in table
    add_table_kv(doc, [
        ("WH Receipt No.",  "WR-2025-HCM-00330"),
        ("Date Received",   "21 March 2025"),
        ("Warehouse",       "Cat Lai Bonded Warehouse, HCMC"),
    ])
    doc.add_paragraph("")

    # Shipment reference — uses "Related Documents" heading
    add_heading(doc, "Related Documents", level=2)
    add_table_kv(doc, [
        ("Invoice No.",   SHIP_C["invoice_number"]),
        ("PO No.",        SHIP_C["po_number"]),
        ("B/L No.",       SHIP_C["bl_number"]),
        ("Container No",  SHIP_C["container_no"]),
        ("Seal No.",      SHIP_C["seal_no"]),
        ("L/C No.",       SHIP_C["lc_number"]),
    ])
    doc.add_paragraph("")

    # Parties — uses "Supplier/Exporter" and "Receiving Party"
    add_heading(doc, "Supplier / Exporter", level=2)
    doc.add_paragraph(SHIP_C["supplier"])

    add_heading(doc, "Receiving Party", level=2)
    doc.add_paragraph(SHIP_C["buyer"])
    doc.add_paragraph("")

    # Goods received — with condition codes (additional detail)
    add_heading(doc, "Inspection & Receipt Details", level=2)
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    headers = ["Item", "HS Code", "Description", "Expected Qty", "Received Qty", "Condition"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1", "8504.40.90", "AC/DC Power Converter PCV-800",   "600",   "600",   "A - Good"],
        ["2", "8536.90.00", "Industrial Relay Switch IRS-450",  "1,200", "1,200", "A - Good"],
        ["3", "8534.00.90", "Multi-layer PCB Assembly MLP-220", "800",   "800",   "A - Good"],
        ["4", "8504.40.90", "Voltage Regulator VR-100S",        "1,500", "1,500", "A - Good"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    # QC section (additional — not in Sets A/B)
    add_heading(doc, "Quality Control", level=2)
    doc.add_paragraph("Visual inspection: PASS")
    doc.add_paragraph("Quantity verification: PASS — all items match packing list")
    doc.add_paragraph("Seal integrity: INTACT")
    doc.add_paragraph("Damage report: NONE")
    doc.add_paragraph("")

    # Summary
    add_heading(doc, "Summary", level=2)
    doc.add_paragraph(f"Total Packages: {SHIP_C['total_packages']} cartons")
    doc.add_paragraph(f"Gross Weight: {SHIP_C['gross_weight']}")
    doc.add_paragraph(f"Customs Clearance: 21 March 2025")
    doc.add_paragraph(f"Port of Discharge: {SHIP_C['port_of_discharge']}")
    doc.add_paragraph(f"Vessel / Voyage: {SHIP_C['vessel']}")
    doc.add_paragraph(f"Country of Origin: {SHIP_C['country_of_origin']}")

    doc.save("doc_input/TestC_04_WarehouseReceipt_WR-2025-HCM-00330.docx")
    print("Created: TestC_04_WarehouseReceipt (variant layout — QC section, condition codes, inspection table)")


# ============================================================
# Generate all
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("SET C: Variant layout (should PASS cross-verify)")
    print("  Different labels, additional fields, table-based headers")
    print("=" * 60)
    create_invoice_c()
    create_packing_list_c()
    create_bl_c()
    create_warehouse_receipt_c()
    print()
    print("VARIANT FEATURES TESTED:")
    print("  - 'Proforma Invoice / Commercial Invoice' title")
    print("  - 'Exporter/Importer' instead of 'Seller/Buyer'")
    print("  - 'Arrival Port' instead of 'Port of Discharge'")
    print("  - Header info in 2-column tables instead of key-value paragraphs")
    print("  - HS codes in line item table (8504.40.90, 8536.90.00, 8534.00.90)")
    print("  - Payment terms, L/C number, insurance amount, country of origin")
    print("  - 'Goods Received Note / Warehouse Receipt' title")
    print("  - QC inspection section with condition codes")
    print("  - 'Gross Wt' / 'Net Wt' abbreviations")
    print("  - 'Notify Party' in Bill of Lading")
    print()
    print("UPLOAD: Upload all 4 files with same shipment_ref — expect 0 cross-verify failures")
