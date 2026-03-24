"""Create test DOCX documents - Set D: VIETNAMESE LANGUAGE.
All labels in Vietnamese. Tests multilingual OCR capability.
Shipment: Hanoi Electronics Manufacturing -> Panasonic Appliances Vietnam
Route: Hai Phong Port -> Ho Chi Minh City Cat Lai Port
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)


def add_table_kv(doc, pairs):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(pairs):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# SHARED SHIPMENT DATA - Set D (Vietnamese)
# ============================================================
SHIP_D = {
    "po_number":         "PO-2025-HN-4400",
    "invoice_number":    "INV-2025-HN-4400",
    "bl_number":         "OOLU7700889900",
    "supplier":          "Hanoi Electronics Manufacturing Co., Ltd.",
    "buyer":             "Panasonic Appliances Vietnam Co., Ltd.",
    "vessel":            "EVERGREEN STAR V.038N",
    "port_of_loading":   "Hai Phong Port",
    "port_of_discharge": "Ho Chi Minh City Cat Lai Port",
    "container_no":      "OOLU6655443",
    "gross_weight":      "18,500.00 kg",
    "net_weight":        "17,200.00 kg",
    "total_packages":    "350",
    "total_amount":      "156,800.00",
    "currency":          "USD",
    "date_invoice":      "20/03/2025",
    "date_bl":           "22/03/2025",
    "date_wr":           "25/03/2025",
    "tax_code":          "0106523478",
    "seal_no":           "SL-20250322-E8",
    "hs_codes":          ["8542.31.00", "8473.30.00", "8504.40.90"],
    "wh_receipt_no":     "WR-2025-HP-00440",
}


# ============================================================
# SET D - Hoa Don Thuong Mai (Commercial Invoice) - Vietnamese
# ============================================================
def create_invoice_d():
    doc = Document()
    add_heading(doc, "HOA DON THUONG MAI", level=1)
    doc.add_paragraph("")

    add_table_kv(doc, [
        ("So hoa don",           SHIP_D["invoice_number"]),
        ("Ngay",                 SHIP_D["date_invoice"]),
        ("So don dat hang",      SHIP_D["po_number"]),
        ("Dong tien",            SHIP_D["currency"]),
        ("Dieu khoan giao hang", "CIF Ho Chi Minh"),
        ("Xuat xu hang hoa",    "Vietnam"),
    ])
    doc.add_paragraph("")

    add_heading(doc, "Ben ban (Nguoi xuat khau)", level=2)
    doc.add_paragraph(SHIP_D["supplier"])
    doc.add_paragraph("So 15, Khu Cong Nghiep Thang Long, Dong Anh, Ha Noi, Viet Nam")
    doc.add_paragraph(f"Ma so thue (MST): {SHIP_D['tax_code']}")

    add_heading(doc, "Ben mua (Nguoi nhap khau)", level=2)
    doc.add_paragraph(SHIP_D["buyer"])
    doc.add_paragraph("Lo CN2, Khu Cong Nghiep Thang Long II, Ha Noi, Viet Nam")
    doc.add_paragraph("")

    add_heading(doc, "Thong tin van chuyen", level=2)
    add_kv(doc, "Tau / Chuyen", SHIP_D["vessel"])
    add_kv(doc, "Cang xuat", SHIP_D["port_of_loading"])
    add_kv(doc, "Cang nhap", SHIP_D["port_of_discharge"])
    add_kv(doc, "So van don", SHIP_D["bl_number"])
    add_kv(doc, "So container", SHIP_D["container_no"])
    add_kv(doc, "So seal", SHIP_D["seal_no"])
    doc.add_paragraph("")

    add_heading(doc, "Chi tiet hang hoa", level=2)
    table = doc.add_table(rows=4, cols=7)
    table.style = 'Table Grid'
    headers = ["STT", "Ma HS", "Mo ta hang hoa", "So luong", "Don vi", "Don gia (USD)", "Thanh tien (USD)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1", "8542.31.00", "Vi mach dieu khien IC-PRO500",    "2,000", "Cai", "42.00", "84,000.00"],
        ["2", "8473.30.00", "Bo phan may tinh BPC-300",        "1,500", "Cai", "28.00", "42,000.00"],
        ["3", "8504.40.90", "Bo chuyen doi nguon PSU-120W",    "800",   "Cai", "38.50", "30,800.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    doc.add_paragraph(f"Tong gia tri CIF: USD {SHIP_D['total_amount']}")
    doc.add_paragraph(f"Tong so kien: {SHIP_D['total_packages']} thung")
    doc.add_paragraph(f"Trong luong ca bi: {SHIP_D['gross_weight']}")
    doc.add_paragraph(f"Trong luong tinh: {SHIP_D['net_weight']}")

    doc.save("doc_input/TestD_01_Invoice_INV-2025-HN-4400.docx")
    print("Created: TestD_01_Invoice (Vietnamese - Hoa Don Thuong Mai)")


# ============================================================
# SET D - Phieu Dong Goi (Packing List) - Vietnamese
# ============================================================
def create_packing_list_d():
    doc = Document()
    add_heading(doc, "PHIEU DONG GOI", level=1)
    doc.add_paragraph("")

    add_table_kv(doc, [
        ("So phieu dong goi",   "PL-2025-HN-4400"),
        ("Ngay",                SHIP_D["date_invoice"]),
        ("Hoa don tham chieu",  SHIP_D["invoice_number"]),
        ("So don dat hang",     SHIP_D["po_number"]),
        ("Xuat xu hang hoa",   "Vietnam"),
    ])
    doc.add_paragraph("")

    add_heading(doc, "Nguoi gui hang", level=2)
    doc.add_paragraph(SHIP_D["supplier"])
    add_heading(doc, "Nguoi nhan hang", level=2)
    doc.add_paragraph(SHIP_D["buyer"])
    doc.add_paragraph("")

    add_heading(doc, "Chi tiet dong goi", level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    headers = ["So thung", "Ma HS", "Mo ta", "SL/Thung", "T.L tinh (kg)", "T.L ca bi (kg)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["001-140", "8542.31.00", "Vi mach dieu khien IC-PRO500",  "15", "5,880.00",  "6,300.00"],
        ["141-260", "8473.30.00", "Bo phan may tinh BPC-300",      "12", "6,480.00",  "7,200.00"],
        ["261-350", "8504.40.90", "Bo chuyen doi nguon PSU-120W",  "9",  "4,840.00",  "5,000.00"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    add_heading(doc, "Tong hop", level=2)
    doc.add_paragraph(f"Tong so kien: {SHIP_D['total_packages']} thung")
    doc.add_paragraph(f"Trong luong ca bi: {SHIP_D['gross_weight']}")
    doc.add_paragraph(f"Trong luong tinh: {SHIP_D['net_weight']}")
    doc.add_paragraph("")

    add_heading(doc, "Ky hieu van chuyen", level=2)
    add_kv(doc, "So container", SHIP_D["container_no"])
    add_kv(doc, "So seal", SHIP_D["seal_no"])
    add_kv(doc, "Tau / Chuyen", SHIP_D["vessel"])
    add_kv(doc, "Cang xuat", SHIP_D["port_of_loading"])
    add_kv(doc, "Cang nhap", SHIP_D["port_of_discharge"])

    doc.save("doc_input/TestD_02_PackingList_PL-2025-HN-4400.docx")
    print("Created: TestD_02_PackingList (Vietnamese - Phieu Dong Goi)")


# ============================================================
# SET D - Van Don Duong Bien (Bill of Lading) - Vietnamese
# ============================================================
def create_bl_d():
    doc = Document()
    add_heading(doc, "VAN DON DUONG BIEN", level=1)
    doc.add_paragraph("")

    add_table_kv(doc, [
        ("So van don",      SHIP_D["bl_number"]),
        ("Ngay phat hanh",  SHIP_D["date_bl"]),
        ("Noi phat hanh",   "Hai Phong, Viet Nam"),
    ])
    doc.add_paragraph("")

    add_heading(doc, "Nguoi gui hang", level=2)
    doc.add_paragraph(SHIP_D["supplier"])
    doc.add_paragraph("So 15, Khu Cong Nghiep Thang Long, Dong Anh, Ha Noi")

    add_heading(doc, "Nguoi nhan hang", level=2)
    doc.add_paragraph(SHIP_D["buyer"])
    doc.add_paragraph("Lo CN2, Khu Cong Nghiep Thang Long II, Ha Noi")

    add_heading(doc, "Ben thong bao", level=2)
    doc.add_paragraph("Panasonic Logistics Vietnam")
    doc.add_paragraph("DT: +84-24-3952-xxxx")
    doc.add_paragraph("")

    add_heading(doc, "Thong tin tau va hanh trinh", level=2)
    add_kv(doc, "Tau / Chuyen", SHIP_D["vessel"])
    add_kv(doc, "Cang xuat", SHIP_D["port_of_loading"])
    add_kv(doc, "Cang nhap", SHIP_D["port_of_discharge"])
    add_kv(doc, "Ngay khoi hanh du kien", "22/03/2025")
    add_kv(doc, "Ngay den du kien", "25/03/2025")
    doc.add_paragraph("")

    add_heading(doc, "Chi tiet hang hoa", level=2)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    headers = ["Container / Seal", "So kien", "Mo ta hang hoa", "Trong luong ca bi", "The tich"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = f"{SHIP_D['container_no']} / {SHIP_D['seal_no']}"
    table.rows[1].cells[1].text = f"{SHIP_D['total_packages']} thung"
    table.rows[1].cells[2].text = ("Linh kien dien tu:\n"
                                    "- Vi mach dieu khien\n"
                                    "- Bo phan may tinh\n"
                                    "- Bo chuyen doi nguon\n"
                                    f"Ma HS: {', '.join(SHIP_D['hs_codes'])}")
    table.rows[1].cells[3].text = SHIP_D["gross_weight"]
    table.rows[1].cells[4].text = "52.40 CBM"
    doc.add_paragraph("")

    doc.add_paragraph(f"Trong luong tinh: {SHIP_D['net_weight']}")
    doc.add_paragraph(f"So hoa don: {SHIP_D['invoice_number']}")
    doc.add_paragraph(f"So don dat hang: {SHIP_D['po_number']}")
    doc.add_paragraph(f"Cuoc van chuyen: USD 3,800.00 (Da tra)")

    doc.save("doc_input/TestD_03_BillOfLading_OOLU7700889900.docx")
    print("Created: TestD_03_BillOfLading (Vietnamese - Van Don Duong Bien)")


# ============================================================
# SET D - Phieu Nhap Kho (Warehouse Receipt) - Vietnamese
# ============================================================
def create_warehouse_receipt_d():
    doc = Document()
    add_heading(doc, "PHIEU NHAP KHO", level=1)
    doc.add_paragraph("")

    add_table_kv(doc, [
        ("So phieu nhap kho",  SHIP_D["wh_receipt_no"]),
        ("Ngay nhap kho",      SHIP_D["date_wr"]),
        ("Kho hang",           "Kho Ngoai Quan Cat Lai, TP.HCM"),
    ])
    doc.add_paragraph("")

    add_heading(doc, "Chung tu lien quan", level=2)
    add_table_kv(doc, [
        ("So hoa don",      SHIP_D["invoice_number"]),
        ("So don dat hang", SHIP_D["po_number"]),
        ("So van don",      SHIP_D["bl_number"]),
        ("So container",    SHIP_D["container_no"]),
        ("So seal",         SHIP_D["seal_no"]),
    ])
    doc.add_paragraph("")

    add_heading(doc, "Nha cung cap", level=2)
    doc.add_paragraph(SHIP_D["supplier"])

    add_heading(doc, "Ben nhan hang", level=2)
    doc.add_paragraph(SHIP_D["buyer"])
    doc.add_paragraph("")

    add_heading(doc, "Chi tiet kiem tra va nhap kho", level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    headers = ["STT", "Ma HS", "Mo ta", "SL du kien", "SL thuc nhan", "Tinh trang"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    items = [
        ["1", "8542.31.00", "Vi mach dieu khien IC-PRO500",  "2,000", "2,000", "Tot"],
        ["2", "8473.30.00", "Bo phan may tinh BPC-300",      "1,500", "1,500", "Tot"],
        ["3", "8504.40.90", "Bo chuyen doi nguon PSU-120W",  "800",   "800",   "Tot"],
    ]
    for r, item in enumerate(items, 1):
        for c, val in enumerate(item):
            table.rows[r].cells[c].text = val
    doc.add_paragraph("")

    add_heading(doc, "Kiem tra chat luong", level=2)
    doc.add_paragraph("Kiem tra ngoai quan: DAT")
    doc.add_paragraph("Xac nhan so luong: DAT - Khop voi phieu dong goi")
    doc.add_paragraph("Tinh trang seal: NGUYEN VEN")
    doc.add_paragraph("Bao cao hu hong: KHONG CO")
    doc.add_paragraph("")

    add_heading(doc, "Tong hop", level=2)
    doc.add_paragraph(f"Tong so kien: {SHIP_D['total_packages']} thung")
    doc.add_paragraph(f"Trong luong ca bi: {SHIP_D['gross_weight']}")
    doc.add_paragraph(f"Thong quan: 25/03/2025")
    doc.add_paragraph(f"Cang nhap: {SHIP_D['port_of_discharge']}")
    doc.add_paragraph(f"Tau / Chuyen: {SHIP_D['vessel']}")
    doc.add_paragraph(f"Xuat xu hang hoa: Vietnam")

    doc.save("doc_input/TestD_04_WarehouseReceipt_WR-2025-HP-00440.docx")
    print("Created: TestD_04_WarehouseReceipt (Vietnamese - Phieu Nhap Kho)")


# ============================================================
# Generate all
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("SET D: Vietnamese language documents")
    print("  All labels in Vietnamese, tests multilingual OCR")
    print("  Hanoi Electronics Mfg -> Panasonic Appliances Vietnam")
    print("  Route: Hai Phong -> Ho Chi Minh City Cat Lai")
    print("=" * 60)
    create_invoice_d()
    create_packing_list_d()
    create_bl_d()
    create_warehouse_receipt_d()
    print()
    print("VIETNAMESE LABELS TESTED:")
    print("  - Hoa Don Thuong Mai (Invoice)")
    print("  - Phieu Dong Goi (Packing List)")
    print("  - Van Don Duong Bien (Bill of Lading)")
    print("  - Phieu Nhap Kho (Warehouse Receipt)")
    print("  - So hoa don, Ngay, Cang xuat/nhap, Tau/Chuyen")
    print("  - Trong luong ca bi/tinh, Tong so kien")
    print("  - Ben ban/mua, Nguoi gui/nhan hang")
    print()
    print("UPLOAD: Upload all 4 files with same shipment_ref")
