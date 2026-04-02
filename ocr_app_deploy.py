import os, io, json, logging, re, traceback, base64
from flask import Flask, request, jsonify
from flask_cors import CORS
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes, convert_from_path

try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
    logging.info("pillow-heif registered")
except ImportError:
    logging.warning("pillow-heif not available")

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import boto3
    BEDROCK_CLIENT = boto3.client("bedrock-runtime", region_name="us-east-1")
    HAS_BEDROCK = True
    logging.info("Bedrock client initialized")
except Exception:
    BEDROCK_CLIENT = None
    HAS_BEDROCK = False
    logging.warning("Bedrock not available")

logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')
log = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
TESSDATA_DIR = os.getenv('TESSDATA_PREFIX', '/usr/local/share/tessdata')
try:
    SUPPORTED_LANGS = [f.replace('.traineddata','') for f in os.listdir(TESSDATA_DIR) if f.endswith('.traineddata')]
except:
    SUPPORTED_LANGS = ['eng', 'vie']

CLAUDE_MODEL_ID = "us.anthropic.claude-sonnet-4-6"

# ---------------------------------------------------------------------------
# Field extraction via regex (for Tesseract OCR text and DOCX text)
# ---------------------------------------------------------------------------

def extract_fields_invoice(text, filename=''):
    """Extract fields from Invoice & Packing List (Panasonic format)."""
    fields = {}

    # Invoice number: MSV250374473 or INV-XXXX
    m = re.search(r'Invoice\s*No\.?\s*[:\s]*\n?\s*([A-Z0-9][\w\-]{6,30})', text, re.IGNORECASE)
    if m:
        fields["invoice_number"] = {"value": m.group(1).strip(), "confidence": 92}
    else:
        m2 = re.search(r'\b(MSV\d{8,12})\b', text)
        if m2:
            fields["invoice_number"] = {"value": m2.group(1).strip(), "confidence": 90}
        else:
            m3 = re.search(r'\b(INV[\-]\d{4}[\-\s][A-Z0-9\-]+)\b', text)
            if m3:
                fields["invoice_number"] = {"value": re.sub(r'\s+', '-', m3.group(1).strip()), "confidence": 85}

    # Date: 25.03.2025 or 18 March 2025
    m = re.search(r'(?:Date|Invoice\s+Date)[:\s]*\n?\s*(\d{1,2}[\./\-]\d{1,2}[\./\-]\d{2,4})', text, re.IGNORECASE)
    if m:
        fields["date"] = {"value": m.group(1).strip(), "confidence": 90}
    else:
        m2 = re.search(r'(?:Date)[:\s]*\n?\s*(\d{1,2}\s+\w+\s+\d{4})', text, re.IGNORECASE)
        if m2:
            fields["date"] = {"value": m2.group(1).strip(), "confidence": 88}
        else:
            m3 = re.search(r'\b(\d{1,2}[\./]\d{1,2}[\./]\d{4})\b', text)
            if m3:
                fields["date"] = {"value": m3.group(1).strip(), "confidence": 75}

    # B/L number from invoice: EGLV142551220956
    m = re.search(r'(?:Bill\s+of\s+Lading|B/?L\s*#?|BL#?|Way\s*Bill\s*No\.?)[:\s]*\n?\s*([A-Z]{4}\d{8,15})', text, re.IGNORECASE)
    if m:
        fields["bl_number"] = {"value": m.group(1).strip(), "confidence": 90}
    else:
        m2 = re.search(r'\b(EGLV\d{10,15})\b', text)
        if not m2:
            m2 = re.search(r'\b([A-Z]{4}\d{10,15})\b', text)
        if m2:
            fields["bl_number"] = {"value": m2.group(1).strip(), "confidence": 85}

    # Shipping reference
    m = re.search(r'Shipping\s+Refer\s*No[:\s]*\n?\s*([A-Z]{4}\s*\d{8,15})', text, re.IGNORECASE)
    if m:
        fields["shipping_ref"] = {"value": m.group(1).strip(), "confidence": 85}

    # Order number
    m = re.search(r'Order\s*No[:\s]*\n?\s*(\d{8,15})', text, re.IGNORECASE)
    if m:
        fields["order_number"] = {"value": m.group(1).strip(), "confidence": 88}

    # Buyer / consignee
    m = re.search(r'(PANASONIC\s+VIETNAM\s+CO\.,?\s*LTD\.?)', text, re.IGNORECASE)
    if m:
        fields["buyer_name"] = {"value": m.group(1).strip(), "confidence": 92}

    # Supplier / shipper
    m = re.search(r'(PANASONIC\s+CONSUMER\s+MARKETING\s+ASIAPACIFIC)', text, re.IGNORECASE)
    if m:
        fields["supplier_name"] = {"value": m.group(1).strip(), "confidence": 90}
    else:
        # Fallback: city-prefix company
        m2 = re.search(r'((?:Shenzhen|Shanghai|Beijing|Guangzhou|Dongguan|Foshan|Ningbo|Xiamen|Hangzhou)\s+[\w\s,\.]+?Co\.[\.,]?\s*Ltd\.?)', text, re.IGNORECASE | re.MULTILINE)
        if m2:
            val = re.sub(r'Co\.[\.,]*\s*Ltd', 'Co., Ltd', m2.group(1).strip())
            fields["supplier_name"] = {"value": val, "confidence": 85}

    # Vessel
    m = re.search(r'(?:Shipped\s+per|Vessel)[:\s]*\n?\s*([A-Z][\w\s]+?(?:\d{4}[\-]\d{2,4}[A-Z]?))', text, re.IGNORECASE)
    if m:
        val = re.sub(r'\s+', ' ', m.group(1).strip())
        # Clean trailing junk
        val = re.sub(r'\s*(PANASONIC|From|ETD).*', '', val).strip()
        if len(val) > 3:
            fields["vessel"] = {"value": val, "confidence": 85}
    if "vessel" not in fields:
        m2 = re.search(r'(?:Vessel\s*/?\s*Voyage)[:\s]*\n?\s*(.+?)(?:\s+Country|\s*$)', text, re.IGNORECASE)
        if m2:
            fields["vessel"] = {"value": m2.group(1).strip(), "confidence": 85}

    # Port of loading / destination
    m = re.search(r'(?:From)[:\s]*\n?\s*([A-Z][A-Z\s]{2,30}?)(?:\s+ETD|\s*$)', text, re.IGNORECASE | re.MULTILINE)
    if m:
        fields["port_of_loading"] = {"value": m.group(1).strip(), "confidence": 85}
    else:
        m2 = re.search(r'(?:Port\s+of\s+Loading)[:\s]*\n?\s*([A-Za-z][\w\s,\.]{3,50})', text, re.IGNORECASE)
        if m2:
            fields["port_of_loading"] = {"value": m2.group(1).strip(), "confidence": 90}

    m = re.search(r'(?:Destination)[:\s]*\n?\s*([A-Z][A-Z\s\.]{2,40}?)(?:\s*$|\n)', text, re.IGNORECASE | re.MULTILINE)
    if m:
        val = m.group(1).strip().rstrip('.')
        if len(val) > 2:
            fields["port_of_discharge"] = {"value": val, "confidence": 85}
    if "port_of_discharge" not in fields:
        m2 = re.search(r'(?:To)[:\s]*\n?\s*([A-Z][A-Z\s]{2,30}?)(?:\s+ETA|\s+Via|\s*$)', text, re.IGNORECASE | re.MULTILINE)
        if m2:
            val = m2.group(1).strip()
            if len(val) > 2:
                fields["port_of_discharge"] = {"value": val, "confidence": 80}
    if "port_of_discharge" not in fields:
        m3 = re.search(r'(?:Port\s+of\s+Discharge|Arrival\s+Port)[:\s]*\n?\s*([A-Za-z][\w\s,\.]{3,50})', text, re.IGNORECASE)
        if m3:
            fields["port_of_discharge"] = {"value": m3.group(1).strip(), "confidence": 90}

    # ETD / ETA
    for tag in ['ETD', 'ETA']:
        m = re.search(rf'{tag}[:\s]*\n?\s*(\d{{1,2}}[\./\-]\d{{1,2}}[\./\-]\d{{2,4}})', text, re.IGNORECASE)
        if m:
            fields[tag.lower()] = {"value": m.group(1).strip(), "confidence": 85}
        else:
            m2 = re.search(rf'{tag}[:\s]*\n?\s*(\d{{1,2}}\s+\w+\s+\d{{4}})', text, re.IGNORECASE)
            if m2:
                fields[tag.lower()] = {"value": m2.group(1).strip(), "confidence": 85}

    # Container numbers: EGHU8301575, EGHU9312640 or MAEU7788001
    containers = list(dict.fromkeys(re.findall(r'\b([A-Z]{4}\d{7})\b', text)))
    if containers:
        fields["container_no"] = {"value": ", ".join(containers[:6]), "confidence": 90}

    # Total amount
    m = re.search(r'(?:TOTAL\s+AMOUNT|Grand\s+Total|TOTAL\s+CIF\s+VALUE)[^:]*?(?:USD\s*)?([\d,]+\.?\d{0,2})', text, re.IGNORECASE)
    if m:
        fields["total_amount"] = {"value": m.group(1).strip(), "confidence": 88}

    # Currency
    m = re.search(r'\b(USD|EUR|JPY|VND|SGD)\b', text)
    if m:
        fields["currency"] = {"value": m.group(1), "confidence": 85}

    # Gross weight
    m = re.search(r'(?:Grs?\s*Wt|Gross\s*W(?:t|eight)?)[^:]*?[:\s]*\n?\s*([\d,]+\.?\d{3})\s*(?:KGS?|kg)', text, re.IGNORECASE)
    if m:
        fields["gross_weight"] = {"value": m.group(1).strip() + " KGS", "confidence": 85}
    else:
        # Fallback: find largest KGS value (must be > 100 to avoid product spec like ">10KG")
        all_wt = re.findall(r'([\d,]+\.?\d{3})\s*(?:KGS|kg)', text, re.IGNORECASE)
        if all_wt:
            candidates = [w for w in all_wt if float(w.replace(',', '')) > 100]
            if candidates:
                best = max(candidates, key=lambda x: float(x.replace(',', '')))
                fields["gross_weight"] = {"value": best + " KGS", "confidence": 75}

    # Net weight
    m = re.search(r'(?:Net\s*W(?:t|eight)?|N\.?W\.?)[^:]*?[:\s]*\n?\s*([\d,]+\.?\d*)\s*(?:KGS?|kg)', text, re.IGNORECASE)
    if m:
        fields["net_weight"] = {"value": m.group(1).strip() + " KGS", "confidence": 85}

    # Measurement / CBM
    m = re.search(r'(?:Measmt|Measurement|CBM)[^:]*?[:\s]*\n?\s*([\d,]+\.?\d*)\s*(?:CBM|CEM|M3|m3)?', text, re.IGNORECASE)
    if m:
        val = m.group(1).strip()
        if float(val.replace(',', '')) > 0:
            fields["measurement"] = {"value": val + " CBM", "confidence": 85}

    # Total packages
    m = re.search(r'(?:Total\s+Packages|No\.?\s+of\s+Packages)[:\s]*\n?\s*(\d[\d,]*)', text, re.IGNORECASE)
    if m:
        fields["total_packages"] = {"value": m.group(1).strip(), "confidence": 85}
    else:
        m2 = re.search(r'(\d[\d,]*)\s+(?:cartons?|ctns?|packages?|sets?)\b', text, re.IGNORECASE)
        if m2:
            fields["total_packages"] = {"value": m2.group(1).strip(), "confidence": 75}

    # Description of goods
    m = re.search(r'(?:WASHING\s+MACHINE|ELECTRONIC\s+COMPONENT|POWER\s+CONVERTER)[\w\s\(\)>\-]*', text, re.IGNORECASE)
    if m:
        fields["description_of_goods"] = {"value": m.group(0).strip()[:100], "confidence": 85}
    if "description_of_goods" not in fields:
        m2 = re.search(r'(?:Description\s+of\s+Goods)[^:]*?[:\s]*\n\s*([A-Z][\w\s\(\)>\-]+)', text, re.IGNORECASE)
        if m2:
            val = m2.group(1).strip()[:100]
            if len(val) > 5 and 'Unit Price' not in val and 'Measmt' not in val:
                fields["description_of_goods"] = {"value": val, "confidence": 80}

    # HS codes
    hs_all = list(set(re.findall(r'\b(\d{4}\.\d{2}(?:\.\d{2})?)\b', text)))
    if hs_all:
        fields["hs_codes"] = {"value": ", ".join(sorted(hs_all)), "confidence": 90}

    # Payment terms
    m = re.search(r'(?:Payment|Pavment)[:\s]*\n?\s*(.+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        fields["payment_terms"] = {"value": m.group(1).strip()[:60], "confidence": 80}

    # PO number
    m = re.search(r'(?:PO\s*(?:No\.?|Ref)?|Purchase\s+Order)[:\s]*\n?\s*(PO[\-][A-Z0-9\-/]{6,30})', text, re.IGNORECASE)
    if m:
        fields["po_number"] = {"value": m.group(1).strip(), "confidence": 88}

    # Incoterms
    m = re.search(r'(?:Incoterms)[:\s]*\n?\s*([A-Z]{3}\s+[\w ]+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        fields["incoterms"] = {"value": m.group(1).strip(), "confidence": 85}
    else:
        m2 = re.search(r'\b(CIF|FOB|CFR|EXW|DDP|DAP)\b', text)
        if m2:
            fields["incoterms"] = {"value": m2.group(1), "confidence": 70}

    # Seal numbers
    seals = re.findall(r'(?:Seal\s*No\.?|SEAL)[:\s]*\n?\s*([A-Z0-9\-]+)', text, re.IGNORECASE)
    if not seals:
        seals = re.findall(r'\b(EMCU[A-Z]{2}\d{4})\b', text)
    if seals:
        fields["seal_no"] = {"value": ", ".join(dict.fromkeys(seals)), "confidence": 85}

    # L/C number
    m = re.search(r'L/C\s*No\.?\s*[:\s]*\n?\s*([A-Z0-9\-]{8,30})', text, re.IGNORECASE)
    if m:
        fields["lc_number"] = {"value": m.group(1).strip(), "confidence": 88}

    # Insurance
    m = re.search(r'Insurance(?:\s+Amount)?[:\s]*\n?\s*(?:USD\s*)?([\d,]+\.?\d{0,2})', text, re.IGNORECASE)
    if m:
        fields["insurance_amount"] = {"value": m.group(1).strip(), "confidence": 85}

    # Country of origin
    m = re.search(r'(?:Country\s+of\s+Origin|MADE\s+IN)\s*[:\s]*\n?\s*([A-Z][A-Za-z ]{1,30})', text, re.IGNORECASE)
    if m:
        fields["country_of_origin"] = {"value": m.group(1).strip(), "confidence": 85}

    # Tax code
    m = re.search(r'(?:Tax\s*Code|MST)[^:]*[:\s]+(\d{10,13})', text, re.IGNORECASE)
    if m:
        fields["tax_code"] = {"value": m.group(1).strip(), "confidence": 90}

    return fields


def extract_fields_co(text, filename=''):
    """Extract fields from Certificate of Origin (Form E)."""
    fields = {}

    # Reference number
    m = re.search(r'Reference\s*No\.?\s*[:\s]*\n?\s*([A-Z0-9\[\]]{10,30})', text, re.IGNORECASE)
    if m:
        val = m.group(1).replace('[', '').replace(']', '').strip()
        fields["co_reference_no"] = {"value": val, "confidence": 88}

    # Exporter
    m = re.search(r'(?:consigned\s+from|Exporter)[^)]*\)\s*\n?\s*([A-Z][\w\s\(\),\.]+?(?:CO\.\s*,?\s*LTD\.?|ASIAPACIFIC))', text, re.IGNORECASE)
    if m:
        fields["supplier_name"] = {"value": m.group(1).strip(), "confidence": 88}

    # Consignee
    m = re.search(r'(PANASONIC\s+VIETNAM\s+CO\.\s*,?\s*LTD\.?)', text, re.IGNORECASE)
    if m:
        fields["buyer_name"] = {"value": m.group(1).strip(), "confidence": 88}
    if "buyer_name" not in fields:
        m2 = re.search(r'(?:consigned\s+to|Consignee)[^)]*\)\s*\n?\s*([A-Z][\w\s\(\),\.]+?(?:CO\.\s*,?\s*LTD\.?|VIETNAM))', text, re.IGNORECASE)
        if m2:
            fields["buyer_name"] = {"value": m2.group(1).strip(), "confidence": 85}

    # Vessel
    m = re.search(r"(?:Vessel'?s?\s*name|Aircraft)[^:]*?[:\s]*\n?\s*([A-Z][\w\s\-]+?\d{3,}[\-\w]*)", text, re.IGNORECASE)
    if m:
        fields["vessel"] = {"value": m.group(1).strip(), "confidence": 85}
    if "vessel" not in fields:
        m2 = re.search(r'\b(EVER\s+CONFORM\s+\d{4}[\-\w]*)\b', text, re.IGNORECASE)
        if m2:
            fields["vessel"] = {"value": m2.group(1).strip(), "confidence": 90}

    # Departure date
    m = re.search(r'Departure\s+date\s*[:\s]*\n?\s*([A-Z]{3}\.?\s*\d{1,2},?\s*\d{4})', text, re.IGNORECASE)
    if m:
        fields["date"] = {"value": m.group(1).strip(), "confidence": 88}

    # Port of discharge
    m = re.search(r'Port\s+of\s+Discharge\s*[:\s]*\n?\s*([A-Z][\w\s,]{3,30}?)(?:\n|FROM|$)', text, re.IGNORECASE)
    if m:
        fields["port_of_discharge"] = {"value": m.group(1).strip().rstrip(','), "confidence": 88}

    # Route: FROM ... TO ...
    m = re.search(r'FROM\s+([A-Z][\w\s,]+?)\s+TO\s+([A-Z][\w\s,]+?)(?:\s+BY|\s*$)', text, re.IGNORECASE | re.MULTILINE)
    if m:
        if "port_of_loading" not in fields:
            fields["port_of_loading"] = {"value": m.group(1).strip(), "confidence": 85}
        if "port_of_discharge" not in fields:
            fields["port_of_discharge"] = {"value": m.group(2).strip(), "confidence": 85}

    # Invoice number and date (field 10)
    m = re.search(r'\b(MSV\d{8,12})\b', text)
    if m:
        fields["invoice_number"] = {"value": m.group(1).strip(), "confidence": 90}
    m = re.search(r'(MAR\.?\s*\d{1,2},?\s*\d{4})', text, re.IGNORECASE)
    if m and "date" not in fields:
        fields["date"] = {"value": m.group(1).strip(), "confidence": 80}

    # HS code
    hs = re.findall(r'(?:HS\s*CODE|IS\s*CODE)[:\s]*\n?\s*(\d{4}[\.\s]\d{2})', text, re.IGNORECASE)
    if hs:
        fields["hs_codes"] = {"value": ", ".join(sorted(set(h.replace(' ', '.') for h in hs))), "confidence": 88}

    # Gross weight / quantity
    m = re.search(r'(\d+)\s*SETS', text, re.IGNORECASE)
    if m:
        fields["total_packages"] = {"value": m.group(1).strip(), "confidence": 85}

    # FOB value
    m = re.search(r'USD[:\s]*([\d,]+\.?\s*\d{0,2})', text)
    if m:
        fields["total_amount"] = {"value": m.group(1).replace(' ', '').strip(), "confidence": 85}
        fields["currency"] = {"value": "USD", "confidence": 90}

    # Description
    m = re.search(r'(?:CARTONS?\s+OF\s+)(.+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        fields["description_of_goods"] = {"value": m.group(1).strip()[:80], "confidence": 80}

    # Country of origin
    m = re.search(r'produced\s+in\s*\n?\s*([A-Z]{3,20})\s*\n', text, re.IGNORECASE)
    if m and m.group(1).strip().isalpha():
        fields["country_of_origin"] = {"value": m.group(1).strip(), "confidence": 88}
    if "country_of_origin" not in fields:
        m2 = re.search(r'Issued\s+in\s+(?:THE\s+)?(?:PEOPLE.S\s+REPUBLIC\s+OF\s+)?([A-Z]{3,20})\b', text, re.IGNORECASE)
        if m2:
            fields["country_of_origin"] = {"value": m2.group(1).strip(), "confidence": 80}

    # Third party operator
    m = re.search(r'THIRD[\-\s]PARTY\s+OPERATOR[:\s]*\n?\s*([A-Z][\w\s]+?)(?:\n\d|\n[A-Z]{3})', text, re.IGNORECASE)
    if m:
        fields["third_party_operator"] = {"value": m.group(1).strip(), "confidence": 80}

    return fields


def extract_fields_bl_claude(raw_bytes):
    """Extract B/L fields using Claude Sonnet 4.6 vision (for image-based B/L forms)."""
    if not HAS_BEDROCK:
        return None, "Bedrock not available"
    try:
        images = convert_from_bytes(raw_bytes, dpi=200)
        if not images:
            return None, "Could not convert PDF to images"

        img = images[0]
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        img_b64 = base64.b64encode(buf.getvalue()).decode()

        prompt = """Extract ALL fields from this Bill of Lading / Sea Waybill document image.
Return a JSON object with these exact keys (use null if not found):
- bl_number: the B/L or document number (digits only, no prefix like EGLV)
- shipper_name: shipper/exporter company name
- consignee_name: consignee full name and address
- notify_party: notify party name and address
- vessel: vessel name and voyage number
- port_of_loading: port of loading city
- port_of_discharge: port of discharge city
- container_no: all container numbers comma separated (format: XXXX1234567)
- seal_no: all seal numbers comma separated
- description_of_goods: goods description
- gross_weight: total gross weight with unit (KGS)
- net_weight: total net weight with unit if available
- measurement: total measurement in CBM
- total_packages: number and type of packages
- freight_terms: PREPAID or COLLECT
- date_of_issue: date of issue
- place_of_issue: place of issue
- export_references: export reference number
- service_type: service type/mode (e.g. FCL/FCL)
- bl_type: SEA WAYBILL or BILL OF LADING

Return ONLY valid JSON, no markdown fences, no explanation."""

        body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 2000,
            "messages": [{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                {"type": "text", "text": prompt}
            ]}]
        })

        resp = BEDROCK_CLIENT.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            body=body,
            contentType="application/json"
        )
        result = json.loads(resp["body"].read())
        claude_text = result["content"][0]["text"]

        # Parse JSON from Claude response
        # Strip markdown fences if present
        claude_text = re.sub(r'^```json\s*', '', claude_text.strip())
        claude_text = re.sub(r'\s*```$', '', claude_text.strip())
        data = json.loads(claude_text)

        # Convert to our field format
        fields = {}
        field_map = {
            "bl_number": "bl_number",
            "shipper_name": "supplier_name",
            "consignee_name": "buyer_name",
            "notify_party": "notify_party",
            "vessel": "vessel",
            "port_of_loading": "port_of_loading",
            "port_of_discharge": "port_of_discharge",
            "container_no": "container_no",
            "seal_no": "seal_no",
            "description_of_goods": "description_of_goods",
            "gross_weight": "gross_weight",
            "net_weight": "net_weight",
            "measurement": "measurement",
            "total_packages": "total_packages",
            "freight_terms": "freight_terms",
            "date_of_issue": "date",
            "place_of_issue": "place_of_issue",
            "export_references": "export_references",
            "service_type": "service_type",
            "bl_type": "bl_type",
        }
        for src_key, dst_key in field_map.items():
            val = data.get(src_key)
            if val and str(val).strip() and str(val).lower() != 'null':
                fields[dst_key] = {"value": str(val).strip(), "confidence": 95}

        return fields, None
    except json.JSONDecodeError as e:
        log.error(f"Claude JSON parse error: {e}, raw: {claude_text[:200]}")
        return None, f"Claude response parse error: {e}"
    except Exception as e:
        log.error(f"Claude B/L extraction error: {traceback.format_exc()}")
        return None, f"Claude error: {e}"


def classify_document(text):
    """Score-based document classification."""
    t = text.lower()
    scores = {
        'invoice': 0,
        'packing_list': 0,
        'bill_of_lading': 0,
        'warehouse_receipt': 0,
        'certificate_of_origin': 0,
    }
    # Certificate of Origin
    for k in ['certificate of origin', 'form e', 'preferential tariff', 'acfta', 'asean-china free trade',
              'products consigned from', 'origin criteria', 'certifying authority']:
        if k in t: scores['certificate_of_origin'] += 25
    # Bill of lading / Sea waybill
    for k in ['bill of lading', 'sea waybill', 'b/l no', 'booking ref', 'place of issue',
              'ocean bill', 'document no', 'shipper', 'pre-carriage', 'onward inland']:
        if k in t: scores['bill_of_lading'] += 20
    # Warehouse receipt
    for k in ['warehouse receipt', 'goods received', 'nhap kho', 'wh receipt no',
              'date received', 'goods received note', 'receiving party', 'inspection']:
        if k in t: scores['warehouse_receipt'] += 25
    # Packing list
    for k in ['packing list', 'carton no', 'packing list no', 'carton marking', 'shipping marks']:
        if k in t: scores['packing_list'] += 20
    if 'ctns' in t: scores['packing_list'] += 10
    # Invoice
    for k in ['commercial invoice', 'invoice & packing list', 'invoice no', 'unit price',
              'amount in words', 'subtotal', 'total cif value', 'proforma invoice',
              'total amount', 'end of invoice']:
        if k in t: scores['invoice'] += 20

    best = max(scores, key=scores.get)
    conf = min(95, scores[best] + 40)
    if scores[best] == 0:
        return 'unknown', 40
    return best, conf


def to_rgb(img):
    if img.mode in ('RGBA', 'LA', 'PA'):
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        return bg
    if img.mode != 'RGB':
        return img.convert('RGB')
    return img


def _extract_from_container(container):
    """Extract text from paragraphs and tables in a docx container."""
    text = ''
    for p in container.paragraphs:
        if p.text.strip():
            text += p.text + '\n'
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    text += ct + '\n'
    return text


def extract_text_from_docx(raw_bytes):
    """Extract text from a .docx file including headers and footers."""
    if not HAS_DOCX:
        return None
    try:
        doc = python_docx.Document(io.BytesIO(raw_bytes))
        text = ''
        for section in doc.sections:
            try: text += _extract_from_container(section.header)
            except: pass
            try: text += _extract_from_container(section.footer)
            except: pass
        text += _extract_from_container(doc)
        return text
    except Exception as e:
        log.error(f"DOCX parse error: {e}")
        return None


def is_bl_document(text, filename=''):
    """Detect if document is a Bill of Lading (needs Claude vision)."""
    t = text.lower() if text else ''
    fn = filename.lower() if filename else ''
    # Check filename hints
    if any(x in fn for x in ['bl_', 'bl-', 'bill_of_lading', 'sea_waybill', 'b_l_', 'bol_']):
        return True
    # Check text content
    bl_keywords = ['bill of lading', 'sea waybill', 'document no', 'shipper', 'consignee',
                   'notify party', 'pre-carriage', 'ocean vessel', 'onward inland']
    score = sum(1 for k in bl_keywords if k in t)
    return score >= 3



def process_single_file(file_obj, lang):
    """Process a single uploaded file: classify + extract fields."""
    raw = file_obj.read()
    fname = file_obj.filename.lower()

    # ---- DOCX: direct text extraction ----
    if fname.endswith('.docx') or fname.endswith('.doc'):
        text = extract_text_from_docx(raw)
        if text is None:
            return None, "Could not parse DOCX file. python-docx may not be installed."
        doc_type, type_conf = classify_document(text)
        if doc_type == 'certificate_of_origin':
            fields = extract_fields_co(text, file_obj.filename)
        else:
            fields = extract_fields_invoice(text, file_obj.filename)
        return {
            "document_id": file_obj.filename.rsplit('.', 1)[0],
            "filename": file_obj.filename,
            "pages": 1,
            "language": lang,
            "document_type": doc_type,
            "type_confidence": type_conf,
            "fields": fields,
            "page_details": [{"page": 1, "text": text, "avg_confidence": 99}],
            "full_text": text
        }, None

    # ---- PDF / Image files ----
    images = []
    try:
        if fname.endswith('.pdf'):
            images = [to_rgb(p) for p in convert_from_bytes(raw, dpi=200)]
        else:
            img = Image.open(io.BytesIO(raw))
            images = [to_rgb(img)]
    except Exception as e:
        log.error(f"Decode error for {file_obj.filename}: {traceback.format_exc()}")
        return None, f"Could not decode: {e}"

    # Quick OCR pass for classification and text extraction
    full_text = ""
    pages = []
    for i, img in enumerate(images):
        try:
            cfg3 = f'--oem 3 --psm 3 -l {lang}'
            txt = pytesseract.image_to_string(img, config=cfg3)
            data = pytesseract.image_to_data(img, config=cfg3, output_type=pytesseract.Output.DICT)
            confs = [int(c) for c in data['conf'] if int(c) > 0]
            avg = round(sum(confs)/len(confs), 1) if confs else 0
        except Exception as e:
            log.error(f"OCR error page {i+1}: {traceback.format_exc()}")
            txt = ""
            avg = 0
        full_text += txt + "\n"
        pages.append({"page": i+1, "text": txt, "avg_confidence": avg})

    doc_type, type_conf = classify_document(full_text)

    # Check if this is a B/L that should use Claude vision
    if doc_type == 'bill_of_lading' or is_bl_document(full_text, file_obj.filename):
        if HAS_BEDROCK:
            log.info(f"Using Claude vision for B/L: {file_obj.filename}")
            fields, err = extract_fields_bl_claude(raw)
            if fields is not None:
                return {
                    "document_id": file_obj.filename.rsplit('.', 1)[0],
                    "filename": file_obj.filename,
                    "pages": len(pages),
                    "language": lang,
                    "document_type": "bill_of_lading",
                    "type_confidence": 95,
                    "fields": fields,
                    "page_details": pages,
                    "full_text": full_text
                }, None
            else:
                log.warning(f"Claude B/L extraction failed ({err}), falling back to Tesseract")

    # Use Tesseract-based extraction for invoice, CO, packing list, etc.
    if doc_type == 'certificate_of_origin':
        fields = extract_fields_co(full_text, file_obj.filename)
    else:
        fields = extract_fields_invoice(full_text, file_obj.filename)

    return {
        "document_id": file_obj.filename.rsplit('.', 1)[0],
        "filename": file_obj.filename,
        "pages": len(pages),
        "language": lang,
        "document_type": doc_type,
        "type_confidence": type_conf,
        "fields": fields,
        "page_details": pages,
        "full_text": full_text
    }, None


@app.route('/health')
def health():
    ver = "unknown"
    try:
        ver = str(pytesseract.get_tesseract_version())
    except:
        pass
    return jsonify({
        "status": "ok",
        "tesseract_version": ver,
        "languages": SUPPORTED_LANGS,
        "docx_support": HAS_DOCX,
        "bedrock_support": HAS_BEDROCK,
        "claude_model": CLAUDE_MODEL_ID
    })


@app.route('/extract', methods=['POST'])
def extract():
    lang = request.form.get('lang', 'eng+vie')
    files = request.files.getlist('file')
    if not files or (len(files) == 1 and files[0].filename == ''):
        return jsonify({"error": "No file provided"}), 400

    if len(files) == 1:
        result, err = process_single_file(files[0], lang)
        if err:
            return jsonify({"error": err}), 422
        log.info(f"Processed: {files[0].filename} -> {result['document_type']}")
        return jsonify(result)

    results = []
    for f in files:
        log.info(f"Batch: {f.filename}")
        result, err = process_single_file(f, lang)
        if err:
            results.append({"filename": f.filename, "error": err})
        else:
            results.append(result)
    return jsonify({"batch": True, "count": len(results), "results": results})


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=False)
